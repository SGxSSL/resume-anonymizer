from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from anonymizer import parse_resume_to_json_gemini
from formatter import format_resume_from_json
import os
import uuid
import shutil
import pdfplumber
from docx import Document
from dotenv import load_dotenv
from typing import List
import json
# Load environment variables from .env file
load_dotenv()

app = FastAPI(
    title="Resume Anonymizer API",
    description="Upload multiple resumes (PDF or DOCX), and get anonymized, formatted DOCX files back.",
)

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],  # Frontend URL
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Define directories for uploads, outputs, and templates
UPLOAD_DIR = "uploads"
OUTPUT_DIR = "outputs"
TEMPLATE_DIR = "templates"
# Create directories if they don't exist
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(TEMPLATE_DIR, exist_ok=True) # Ensure template directory exists

# --- Note on Templates ---
# The formatter.py file uses a template: Document("templates/company_template.docx")
# Make sure you have a file named 'company_template.docx' inside a 'templates' folder.
# If you don't have one, create a simple empty .docx file at that location.

@app.post("/anonymize", tags=["Resume Processing"])
async def anonymize_resumes(files: List[UploadFile] = File(..., description="Resume files in .pdf or .docx format")):
    """
    Processes multiple uploaded resume files:
    1. Extracts text content from each file.
    2. Anonymizes the content using the Google Gemini API.
    3. Formats the anonymized data into new DOCX files based on a template.
    4. Returns download URLs for the generated files.
    """
    output_files = []
    
    for file in files:
        file_extension = os.path.splitext(file.filename)[1].lower()
        if file_extension not in [".pdf", ".docx"]:
            raise HTTPException(status_code=400, detail=f"Unsupported file type for {file.filename}. Please upload .pdf or .docx files only.")

        file_id = str(uuid.uuid4())
        input_path = os.path.join(UPLOAD_DIR, f"{file_id}{file_extension}")
        
        try:
            # Save the uploaded file temporarily
            with open(input_path, "wb") as f:
                content = await file.read()
                f.write(content)

            # Extract text from the saved file
            text = ""
            if file_extension == ".pdf":
                try:
                    with pdfplumber.open(input_path) as pdf:
                        text = "\n".join(page.extract_text() for page in pdf.pages if page.extract_text())
                except Exception as e:
                    raise HTTPException(status_code=500, detail=f"Failed to extract text from PDF: {e}")
            elif file_extension == ".docx":
                try:
                    doc = Document(input_path)
                    text = "\n".join(p.text for p in doc.paragraphs)
                except Exception as e:
                    raise HTTPException(status_code=500, detail=f"Failed to extract text from DOCX: {e}")
            
            if not text.strip():
                raise HTTPException(status_code=400, detail=f"Could not extract any text from {file.filename}")

            # Parse with Gemini
            try:
                parsed_data = parse_resume_to_json_gemini(text)
                print(json.dumps(parsed_data, indent=2))  # Ensure JSON is formatted for debugging
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"Failed to parse {file.filename} with the AI model: {e}")

            # Format into new DOCX
            original_filename = os.path.splitext(os.path.basename(file.filename))[0]
            output_filename = f"{file_id}_{original_filename}_anonymized.docx"
            output_path = os.path.join(OUTPUT_DIR, output_filename)
            
            try:
                format_resume_from_json(parsed_data, output_path)
                output_files.append({
                    "originalName": file.filename,
                    "downloadUrl": f"/download/{output_filename}"
                })
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"Failed to format {file.filename}. Error: {e}")

        finally:
            # Clean up the input file
            if os.path.exists(input_path):
                os.remove(input_path)

    return {"downloadUrls": [f"http://localhost:8000{file['downloadUrl']}" for file in output_files]}

@app.get("/download/{filename}")
async def download_file(filename: str):
    """
    Download a processed resume file
    """
    file_path = os.path.join(OUTPUT_DIR, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")
    
    return FileResponse(
        path=file_path,
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )