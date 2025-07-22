# formatter.py
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), "templates")

def format_resume_from_json(data: dict, output_path: str, template_filename: str = "company_template.docx"):
    # Create a new document
    doc = Document()
    
    # Set up page margins for a clean look
    sections = doc.sections
    for section in sections:
        section.left_margin = Pt(72)    # 1 inch
        section.right_margin = Pt(72)   # 1 inch
        section.top_margin = Pt(72)     # 1 inch
        section.bottom_margin = Pt(72)  # 1 inch
    
    # Add the name as a title with extra spacing
    name_paragraph = doc.add_paragraph()
    name_run = name_paragraph.add_run(data.get("Name", ""))
    name_run.bold = True
    name_run.font.size = Pt(20)
    name_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add extra spacing after name
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Add a line break
    doc.add_paragraph()

    if summary := data.get("Summary"):
        heading = doc.add_paragraph()
        heading_run = heading.add_run("Professional Summary")
        heading_run.bold = True
        heading_run.font.size = Pt(16)
        heading_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Add double line under the heading
        heading.paragraph_format.space_after = Pt(0)
        border_paragraph = doc.add_paragraph("_" * 80)
        border_paragraph.paragraph_format.space_after = Pt(12)
        
        doc.add_paragraph(summary)
        doc.add_paragraph()  # Add spacing

    if skills := data.get("Skills"):
        heading = doc.add_paragraph()
        heading_run = heading.add_run("Technical Skills")
        heading_run.bold = True
        heading_run.font.size = Pt(16)
        heading_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Add double line under the heading
        heading.paragraph_format.space_after = Pt(0)
        border_paragraph = doc.add_paragraph("_" * 80)
        border_paragraph.paragraph_format.space_after = Pt(12)
        
        skills_para = doc.add_paragraph()
        skills_para.add_run(", ".join(skills))
        doc.add_paragraph()  # Add spacing

    if exp_list := data.get("Experience"):
        # Professional History section
        heading = doc.add_paragraph()
        heading_run = heading.add_run("Professional History")
        heading_run.bold = True
        heading_run.font.size = Pt(16)
        heading_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Add double line under the heading
        heading.paragraph_format.space_after = Pt(0)
        border_paragraph = doc.add_paragraph("_" * 80)
        border_paragraph.paragraph_format.space_after = Pt(12)
        
        for exp in exp_list:
            # Company and title
            p = doc.add_paragraph()
            title_run = p.add_run(f"{exp['company']}")
            title_run.bold = True
            title_run.font.size = Pt(12)
            
            # Role and dates on next line
            p = doc.add_paragraph()
            role_run = p.add_run(f"{exp['job_title']}")
            role_run.italic = True
            p.add_run(f" ({exp['dates']})")
            
            # Description with bullet points
            desc_lines = exp['description'].split('\n')
            for line in desc_lines:
                if line.strip():
                    bullet_p = doc.add_paragraph(style='List Bullet')
                    bullet_p.add_run(line.strip())
            
            doc.add_paragraph()  # Add spacing

    if edu_list := data.get("Education"):
        heading = doc.add_paragraph()
        heading_run = heading.add_run("Education")
        heading_run.bold = True
        heading_run.font.size = Pt(16)
        heading_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Add double line under the heading
        heading.paragraph_format.space_after = Pt(0)
        border_paragraph = doc.add_paragraph("_" * 80)
        border_paragraph.paragraph_format.space_after = Pt(12)
        
        for edu in edu_list:
            # School name
            p = doc.add_paragraph()
            school_run = p.add_run(f"{edu['school']}")
            school_run.bold = True
            school_run.font.size = Pt(12)
            
            # Degree and dates
            p = doc.add_paragraph()
            degree_run = p.add_run(f"{edu['degree']}")
            degree_run.italic = True
            p.add_run(f" ({edu['dates']})")
            
            if edu.get('description'):
                desc_p = doc.add_paragraph()
                desc_p.add_run(edu['description'])
            
            doc.add_paragraph()  # Add spacing

    if projects := data.get("Projects"):
        heading = doc.add_paragraph()
        heading_run = heading.add_run("Projects")
        heading_run.bold = True
        heading_run.font.size = Pt(16)
        heading_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Add double line under the heading
        heading.paragraph_format.space_after = Pt(0)
        border_paragraph = doc.add_paragraph("_" * 80)
        border_paragraph.paragraph_format.space_after = Pt(12)
        
        for project in projects:
            # Project title
            p = doc.add_paragraph()
            project_run = p.add_run(project['title'])
            project_run.bold = True
            project_run.font.size = Pt(12)
            if project.get('dates'):
                p.add_run(f" ({project['dates']})")
            
            # Technologies
            if tech := project.get('technologies'):
                tech_p = doc.add_paragraph()
                tech_p.add_run("Technologies: ").bold = True
                
                # Handle both string and list formats, and clean up the technology string
                if isinstance(tech, str):
                    # Remove extra spaces and split by commas
                    cleaned_tech = "".join(tech.split())  # Remove all whitespace
                    tech_list = [t.strip() for t in cleaned_tech.split(',') if t.strip()]
                else:
                    tech_list = tech
                
                tech_p.add_run(", ".join(tech_list))
            
            # Description with bullet points
            if desc := project.get('description'):
                desc_lines = desc.split('\n')
                for line in desc_lines:
                    if line.strip():
                        # Keep existing bullet points, add bullets if not present
                        line = line.strip()
                        if not line.startswith('•'):
                            bullet_p = doc.add_paragraph(style='List Bullet')
                        else:
                            bullet_p = doc.add_paragraph()
                        bullet_p.add_run(line)
            
            doc.add_paragraph()  # Add spacing

    if achievements := data.get("Achievements"):
        heading = doc.add_paragraph()
        heading_run = heading.add_run("Achievements")
        heading_run.bold = True
        heading_run.font.size = Pt(14)
        
        # Create a bullet list for achievements
        for achievement in achievements:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(achievement)
        doc.add_paragraph()  # Add spacing

    # Save the document
    doc.save(output_path)
